#!/usr/bin/env python3
"""
Sakura Anki DOCX 导出器
从 logs/YYYY-MM-DD.md 解析 Q&A 对，生成符合 word.docx 排版规范的 DOCX 文件。

用法:
    python export_anki_docx.py --date 2026-03-30
    python export_anki_docx.py --date 2026-03-30 --subject "AP-Calculus-BC"
    python export_anki_docx.py --date 2026-03-30 --path "C:/path/to/"
    python export_anki_docx.py --date 2026-03-30 --output "output/2026-03-30-anki.docx"

环境变量:
    SAKURA_BASE_PATH  覆盖默认设定目录
"""

import argparse
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# 路径配置
# ---------------------------------------------------------------------------
DEFAULT_WINDOWS_BASE = Path(r"C:\Users\25472\Sakura - gemini版")


def resolve_base_path(cli_path: str | None) -> Path:
    if cli_path:
        p = Path(cli_path)
        if not p.exists():
            raise FileNotFoundError(f"指定的路径不存在: {p}")
        return p.resolve()
    env_path = os.environ.get("SAKURA_BASE_PATH")
    if env_path:
        p = Path(env_path)
        if p.exists():
            return p.resolve()
    script_dir = Path(__file__).parent.resolve()
    sample_dir = script_dir / "sample_files"
    if sample_dir.exists():
        return sample_dir
    raise FileNotFoundError(
        "未找到设定目录。请通过 --path 参数、环境变量 SAKURA_BASE_PATH "
        "或确保 sample_files/ 目录存在。"
    )


# ---------------------------------------------------------------------------
# 数据结构
# ---------------------------------------------------------------------------
@dataclass
class QCard:
    """一张 Anki 卡片（来自 logs/ 的一个 [Qn]/[An] 对）。"""
    number: int          # Q编号，如 1, 2, 3
    question: str
    answer: str
    section: str          # 来源章节标题，如 "三笠老师：导数应用"
    subject: str          # 科目标签，如 "AP-Calculus-BC"
    lesson_tag: str       # 课次标签，如 "lesson-0005"
    question_type: str   # 题型：单选/多选/判断/简答/名词解释/填空
    color_hints: list[tuple[str, str]] = field(default_factory=list)
    # 每项为 (color_hex, text) 表示需要着色的一段文字


# ---------------------------------------------------------------------------
# 日志解析
# ---------------------------------------------------------------------------

# 注意: **[Q1]** 是字面量 ** + [Q1] + **，不是正则字符类
# 匹配 **[Q1]** 问题文本（字面 star-star, 括号, star-star）
Q_PATTERN = re.compile(r"\*\*\[Q(\d+)\]\*\*\s+(.+?)(?=\*\*\[(?:Q|A)\d+\]\*\*|\Z)", re.DOTALL)
A_PATTERN = re.compile(r"\*\*\[A(\d+)\]\*\*\s+(.+?)(?=\*\*\[(?:Q|A)\d+\]\*\*|\Z)", re.DOTALL)

# 匹配章节标题
SECTION_PATTERN = re.compile(r"^##\s+(.+)$", re.MULTILINE)


def parse_log_file(content: str) -> list[dict]:
    """
    解析日志文件，返回 [
      {section, qa_pairs: [(number, q_text, a_text), ...]},
      ...
    ]
    """
    result = []
    # 按章节分段
    parts = re.split(r"(?=^##\s+)", content, flags=re.MULTILINE)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        lines = part.splitlines()
        # 第一行是 ## 标题
        if not lines or not lines[0].startswith("##"):
            continue
        section_title = lines[0][2:].strip()  # 去掉 "## "
        body = "\n".join(lines[1:]).strip()

        # 提取所有 [Qn] / [An] 对
        # 先找所有标记位置
        qa_blocks: list[tuple[int, str, str]] = []
        # 找到所有 Q 和 A 及其内容
        q_matches = list(Q_PATTERN.finditer(body))
        a_matches = list(A_PATTERN.finditer(body))

        # 构建 (pos, label, content) 列表
        markers: list[tuple[int, str, str]] = []
        for m in q_matches:
            markers.append((m.start(), "Q", m.group(2).strip()))
        for m in a_matches:
            markers.append((m.start(), "A", m.group(2).strip()))
        markers.sort(key=lambda x: x[0])

        # 配对
        pairs: list[tuple[int, str, str]] = []
        i = 0
        while i < len(markers):
            if markers[i][1] == "Q" and i + 1 < len(markers) and markers[i+1][1] == "A":
                q_num_match = re.search(r"Q(\d+)", body[markers[i][0]:markers[i][0]+10])
                num = int(q_num_match.group(1)) if q_num_match else 0
                pairs.append((num, markers[i][2], markers[i+1][2]))
                i += 2
            else:
                i += 1

        if pairs:
            result.append({"section": section_title, "qa_pairs": pairs})
    return result


# ---------------------------------------------------------------------------
# 题型自动判断
# ---------------------------------------------------------------------------
def infer_question_type(question: str, answer: str) -> str:
    """根据内容推断题型。"""
    q = question
    a = answer
    # 选择题：含 A. B. C. D 或 (A) (B) (C) (D)
    if re.search(r"\b[A-D]\.\s*\w|\([A-D]\)\s*\w", q):
        # 统计选项数
        opts = re.findall(r"\b[A-D]\.|\([A-D]\)", q)
        if len(opts) > 2:
            return "多选题"
        return "单选题"
    # 判断题：含 "对/错" 或 "True/False"
    if re.search(r"[对错]|[Tt]rue\s*,?\s*[Ff]alse", a):
        return "判断题"
    # 填空题：含 "____" 或 "……" 或 "（    ）"
    if re.search(r"_{2,}|……|（\s*）|\（\s*\)", q):
        return "填空题"
    # 名词解释：问"什么是X"且答案较短
    if re.search(r"什么是|什么叫|定义是|名词解释", q) and len(a) < 200:
        return "名词解释"
    # 默认简答题
    return "简答题"


# ---------------------------------------------------------------------------
# 科目推断
# ---------------------------------------------------------------------------
SUBJECT_KEYWORDS = {
    "AP-Calculus-BC": [
        "极限", "连续", "导数", "求导", "隐函数", "积分", "微分",
        "ε", "δ", "chain rule", "derivative", "limit", "洛必达",
        "极值", "临界点", "拐点", "泰勒", "级数", "可导",
    ],
    "Physics": ["牛顿", "力学", "动能", "势能", "动量", "能量"],
}


def infer_subject(section: str, question: str, answer: str) -> str:
    text = section + " " + question + " " + answer
    for subject, keywords in SUBJECT_KEYWORDS.items():
        if any(kw in text for kw in keywords):
            return subject
    return "default"


def infer_lesson_tag(section: str, question: str, answer: str) -> str:
    """从内容推断课次标签（从 progress.md 匹配）。"""
    # 简单 heuristics：section 中的课次
    m = re.search(r"第\s*([0-9]{4})\s*课", section)
    if m:
        return f"lesson-{m.group(1)}"
    # 从内容中搜索
    m = re.search(r"第\s*([0-9]{4})\s*课", question + answer)
    if m:
        return f"lesson-{m.group(1)}"
    return "lesson-0000"


# ---------------------------------------------------------------------------
# DOCX 生成
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] python-docx 未安装。请运行: pip install python-docx")
    sys.exit(1)


HEX_TO_RGB = {
    "C00000": RGBColor(0xC0, 0x00, 0x00),   # 填空挖空（红）
    "E97132": RGBColor(0xE9, 0x71, 0x32),   # 批注文本（橙）
    "00B0F0": RGBColor(0x00, 0xB0, 0xF0),   # 批注内容（蓝）
    "F64926": RGBColor(0xF6, 0x49, 0x26),   # 折叠标题（橙红）
    "0E2841": RGBColor(0x0E, 0x28, 0x41),   # 折叠内容（深蓝）
    "CF4DAA": RGBColor(0xCF, 0x4D, 0xAA),   # 隐藏标签（粉紫）
    "3A7C22": RGBColor(0x3A, 0x7C, 0x22),   # 隐藏内容（绿）
}


def hex_to_rgb(hex_str: str) -> RGBColor:
    """'C00000' → RGBColor."""
    hex_str = hex_str.lstrip("#")
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def set_run_color(run, hex_color: str):
    """设置 run 文字颜色。"""
    run.font.color.rgb = hex_to_rgb(hex_color)


def add_colored_run(para, text: str, hex_color: str, bold: bool = False, italic: bool = False):
    """向段落添加带颜色的文字片段。"""
    run = para.add_run(text)
    run.font.color.rgb = hex_to_rgb(hex_color)
    run.bold = bold
    run.italic = italic
    return run


def make_subdeck_heading(doc: Document, subject: str):
    """添加子牌组标题行 `# AP-Calculus-BC`"""
    p = doc.add_paragraph()
    run = p.add_run(f"# {subject}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x60, 0x9E)  # 深蓝色


def add_tag_line(doc: Document, lesson_tag: str, extra_tags: list[str] = None):
    """添加标签行 `TAG lesson-0005 ...`"""
    tags = [lesson_tag] + (extra_tags or [])
    p = doc.add_paragraph()
    run = p.add_run(f"TAG {', '.join(tags)}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)  # 灰色


def add_qa_block(doc: Document, card: QCard):
    """添加一张 Q&A 卡片块到文档。"""
    # 题型前缀
    p_type = doc.add_paragraph()
    run = p_type.add_run(f"* {card.question_type}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    # 问题
    p_q = doc.add_paragraph()
    run_q = p_q.add_run(card.question)
    run_q.font.size = Pt(11)

    # 答案
    p_a = doc.add_paragraph()
    run_a_label = p_a.add_run("答案：")
    run_a_label.bold = True
    run_a_label.font.size = Pt(11)

    # 根据题型应用颜色
    if card.question_type == "填空题":
        # 答案用红色
        run_a = p_a.add_run(card.answer)
        run_a.font.size = Pt(11)
        run_a.font.color.rgb = hex_to_rgb("C00000")
    elif card.question_type == "名词解释":
        run_a = p_a.add_run(card.answer)
        run_a.font.size = Pt(11)
        run_a.font.color.rgb = RGBColor(0x1A, 0x7A, 0x4A)  # 深绿
    elif card.question_type == "判断题":
        # 判断题答案用蓝色标注
        run_a = p_a.add_run(card.answer)
        run_a.font.size = Pt(11)
        run_a.font.color.rgb = hex_to_rgb("00B0F0")
    else:
        run_a = p_a.add_run(card.answer)
        run_a.font.size = Pt(11)

    # 分隔线
    doc.add_paragraph("─" * 40)


def generate_docx(cards: list[QCard], output_path: Path, subject_filter: Optional[str] = None):
    """生成 DOCX 文件。"""
    if subject_filter:
        cards = [c for c in cards if c.subject == subject_filter]

    if not cards:
        print("[WARN] 没有可导出的卡片。")
        return

    doc = Document()

    # 文档标题
    title = doc.add_heading("Sakura 学习卡片导出", level=1)

    # 按科目分组
    current_subject = None
    current_lesson = None
    for card in cards:
        # 换科目时加标题
        if card.subject != current_subject:
            current_subject = card.subject
            doc.add_heading(f"牌组: {current_subject}", level=2)
        # 换课次时加标签
        if card.lesson_tag != current_lesson:
            current_lesson = card.lesson_tag
            p = doc.add_paragraph()
            run = p.add_run(f"TAG {card.lesson_tag}, {card.subject}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        add_qa_block(doc, card)

    doc.save(str(output_path))
    print(f"[INFO] DOCX 已保存: {output_path}")


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Sakura Anki DOCX 导出器")
    parser.add_argument(
        "--date", "-d",
        required=True,
        help="日志日期，格式 YYYY-MM-DD",
    )
    parser.add_argument(
        "--subject", "-s",
        help="只导出指定科目的卡片 (如 'AP-Calculus-BC')",
    )
    parser.add_argument(
        "--path", "-p",
        help="设定文件所在目录（覆盖默认 Windows 路径）",
    )
    parser.add_argument(
        "--output", "-o",
        help="输出文件路径（默认: output/YYYY-MM-DD-anki.docx）",
    )
    args = parser.parse_args()

    base = resolve_base_path(args.path)
    log_dir = base / "logs"
    log_file = log_dir / f"{args.date}.md"

    print(f"[INFO] 设定目录: {base}")
    print(f"[INFO] 日志文件: {log_file}")

    if not log_file.exists():
        # 尝试 sample_files/logs/
        sample_log = Path(__file__).parent / "sample_files" / "logs" / f"{args.date}.md"
        if sample_log.exists():
            log_file = sample_log
            print(f"[INFO] 使用示例日志: {log_file}")
        else:
            print(f"[ERROR] 日志文件不存在: {log_file}")
            sys.exit(1)

    content = log_file.read_text(encoding="utf-8")
    parsed = parse_log_file(content)

    if not parsed:
        print("[ERROR] 未能从日志中解析出 Q&A 对。")
        sys.exit(1)

    print(f"[INFO] 解析了 {len(parsed)} 个章节")

    # 构建卡片列表
    cards: list[QCard] = []
    for section_data in parsed:
        section_title = section_data["section"]
        qa_pairs = section_data["qa_pairs"]
        for num, q_text, a_text in qa_pairs:
            q_type = infer_question_type(q_text, a_text)
            subject = infer_subject(section_title, q_text, a_text)
            lesson_tag = infer_lesson_tag(section_title, q_text, a_text)
            cards.append(QCard(
                number=num,
                question=q_text.strip(),
                answer=a_text.strip(),
                section=section_title,
                subject=subject,
                lesson_tag=lesson_tag,
                question_type=q_type,
            ))

    total = len(cards)
    print(f"[INFO] 共提取 {total} 张卡片")

    if args.subject:
        subject_filter = args.subject.strip()
        before = len(cards)
        cards = [c for c in cards if c.subject == subject_filter]
        print(f"[INFO] 科目筛选 '{subject_filter}': {before} → {len(cards)} 张卡片")
    else:
        print(f"[INFO] 科目分布:")
        from collections import Counter
        for subj, cnt in Counter(c.subject for c in cards).items():
            print(f"  - {subj}: {cnt} 张")

    if not cards:
        print("[WARN] 筛选后无卡片。")
        sys.exit(0)

    # 输出
    output_path = Path(args.output) if args.output else Path("output")
    output_path.mkdir(exist_ok=True)
    output_file = output_path / f"{args.date}-anki.docx"

    generate_docx(cards, output_file, args.subject)

    print(f"[INFO] 完成！")
    print(f"[INFO] 共导出 {len(cards)} 张卡片")
    print(f"[INFO] Anki 导入: 直接打开 DOCX 或用 '文件→导入' 选择 DOCX 文件")


if __name__ == "__main__":
    main()
